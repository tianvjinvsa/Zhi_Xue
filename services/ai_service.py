"""
AI服务 - 处理AI相关功能
"""
import json
import base64
from pathlib import Path
from typing import List, Optional, Dict
import re

from config import config as app_config
from models import Question

# 可选的文件处理依赖
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from PIL import Image
    import io
    HAS_PIL = True
except ImportError:
    HAS_PIL = False


class AIService:
    """AI服务类"""
    
    # 题目生成的Prompt模板
    QUESTION_PARSE_PROMPT = """你是一个专业的题目格式化助手。请将以下内容转换为标准的JSON格式题目。

要求：
1. 识别题目类型：single(单选题), multiple(多选题), judge(判断题), fill(填空题)
2. 提取题目内容、选项(如有)、正确答案
3. 如果能推断出答案解析，请补充explanation字段
4. 难度默认为3(1-5级)
5. 如果有多道题目，返回JSON数组
6. [重要]题目内容中不要包含原始的题目序号(如"1."、"第1题"、"一、"等)，只保留纯粹的题目内容

输入内容：
{content}

请严格按照以下JSON格式输出(不要添加任何其他文字说明)：
{
  "questions": [
    {
      "type": "题目类型(single/multiple/judge/fill)",
      "question": "题目内容(不含序号，直接是题目正文)",
      "options": ["A. 内容1", "B. 内容2", "C. 内容3", "D. 内容4"],
      "answer": "正确答案(单选为字母如A，多选为数组如[\"A\",\"B\"]，判断为正确/错误)",
      "explanation": "答案解析(可选)",
      "difficulty": 3
    }
  ]
}

注意：
- 选项格式统一为 "X. 内容" 的形式
- 判断题不需要options字段
- 多选题的answer必须是数组格式
- 确保输出是合法的JSON格式
- 题目内容必须去除开头的序号，例如"1. 什么是..."应该变成"什么是..." """

    QUESTION_GENERATE_PROMPT = """你是一个专业的出题助手。请根据以下知识点或主题生成题目。

知识点/主题：
{topic}

要求：
1. 生成 {count} 道题目
2. 题目类型分布：{type_distribution}
3. 难度范围：{difficulty_range}
4. 题目要有区分度，难度适中
5. 选项设计要有干扰性，但答案必须明确

请严格按照以下JSON格式输出：
{
  "questions": [
    {
      "type": "题目类型",
      "question": "题目内容",
      "options": ["A. 内容1", "B. 内容2", "C. 内容3", "D. 内容4"],
      "answer": "正确答案",
      "explanation": "答案解析",
      "difficulty": 难度数字
    }
  ]
}"""

    def __init__(self):
        self._client = None
    
    def _get_client(self):
        """获取或创建API客户端"""
        if self._client is None:
            try:
                from openai import OpenAI
                
                ai_config = app_config.ai_config
                
                kwargs = {
                    'api_key': ai_config.api_key,
                    'timeout': 60.0,  # 设置60秒超时
                }
                
                if ai_config.api_base_url:
                    # 确保 base_url 格式正确
                    base_url = ai_config.api_base_url.strip()
                    # 移除末尾的斜杠
                    if base_url.endswith('/'):
                        base_url = base_url.rstrip('/')
                    # 如果用户误输入了 /chat/completions，将其移除，SDK会自动添加
                    if base_url.endswith('/chat/completions'):
                        base_url = base_url.replace('/chat/completions', '')
                    if base_url.endswith('/chat'):
                        base_url = base_url.replace('/chat', '')
                        
                    kwargs['base_url'] = base_url
                
                self._client = OpenAI(**kwargs)
            except ImportError:
                raise RuntimeError("请安装openai库: pip install openai")
            except Exception as e:
                raise RuntimeError(f"初始化AI客户端失败: {e}")
        
        return self._client
    
    def _reset_client(self):
        """重置客户端(配置更改后调用)"""
        self._client = None
    
    def check_connection(self, temp_config: Optional[Dict] = None) -> tuple[bool, str]:
        """
        检查AI服务连接
        :param temp_config: 临时配置字典，如果提供则使用该配置进行测试
        """
        try:
            # 如果提供了临时配置，创建一个临时的 client
            client_to_use = self._client
            model_to_use = app_config.ai_config.model
            
            if temp_config:
                from openai import OpenAI
                
                api_key = temp_config.get("api_key")
                # 如果临时配置中没有 key 但有 url/model，尝试使用保存的 key
                if not api_key and app_config.ai_config.api_key:
                    api_key = app_config.ai_config.api_key
                    
                if not api_key:
                    return False, "未设置 API Key"
                    
                base_url = temp_config.get("api_base_url") or app_config.ai_config.api_base_url
                model_to_use = temp_config.get("model") or app_config.ai_config.model
                
                try:
                    client_to_use = OpenAI(
                        api_key=api_key,
                        base_url=base_url,
                        timeout=10.0
                    )
                except Exception as e:
                    return False, f"客户端初始化失败: {str(e)}"
            
            # 确保 client 已初始化
            if not client_to_use:
                client_to_use = self._get_client()
                
            if not client_to_use:
                return False, "AI服务未配置"

            # 发送一个简单的请求
            response = client_to_use.chat.completions.create(
                model=model_to_use,
                messages=[
                    {"role": "user", "content": "Hi"}
                ],
                max_tokens=5
            )
            return True, "连接成功"
        except Exception as e:
            return False, f"连接失败: {str(e)}"
    
    def _call_api(self, messages: List[Dict], use_vision: bool = False) -> str:
        """调用API"""
        client = self._get_client()
        ai_config = app_config.ai_config
        
        # 如果是视觉请求但未配置视觉模型，则回退到普通模型
        model = ai_config.model
        if use_vision and ai_config.vision_model:
            model = ai_config.vision_model
        
        print(f"正在调用 AI API, 模型: {model}, Base URL: {ai_config.api_base_url}")
        
        try:
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                max_tokens=ai_config.max_tokens,
                temperature=ai_config.temperature
            )
            content = response.choices[0].message.content
            print(f"AI 响应成功，长度: {len(content)}")
            return content
        except Exception as e:
            print(f"AI 调用失败: {str(e)}")
            raise e
    
    def _parse_json_response(self, response: str) -> Dict:
        """解析JSON响应"""
        # 尝试直接解析
        try:
            return json.loads(response)
        except:
            pass
        
        # 尝试提取JSON块
        json_match = re.search(r'\{[\s\S]*\}', response)
        if json_match:
            try:
                return json.loads(json_match.group())
            except:
                pass
        
        # 尝试提取JSON数组
        json_match = re.search(r'\[[\s\S]*\]', response)
        if json_match:
            try:
                return {"questions": json.loads(json_match.group())}
            except:
                pass
        
        raise ValueError("无法解析AI返回的内容为JSON格式")
    
    def parse_questions_from_text(self, text: str) -> tuple[List[Question], str]:
        """
        从文本解析题目
        返回: (题目列表, 错误消息)
        """
        try:
            # 使用 replace 而不是 format，防止 text 中的花括号导致 KeyError
            prompt = self.QUESTION_PARSE_PROMPT.replace("{content}", text)
            
            messages = [
                {"role": "system", "content": "你是一个专业的题目格式化助手，只输出JSON格式数据。"},
                {"role": "user", "content": prompt}
            ]
            
            response = self._call_api(messages)
            data = self._parse_json_response(response)
            
            questions = []
            for q_data in data.get('questions', []):
                try:
                    q = Question(
                        type=q_data.get('type', 'single'),
                        question=q_data.get('question', ''),
                        options=q_data.get('options', []),
                        answer=q_data.get('answer', ''),
                        explanation=q_data.get('explanation', ''),
                        difficulty=q_data.get('difficulty', 3),
                        source='ai_generated'
                    )
                    questions.append(q)
                except Exception as e:
                    print(f"解析单个题目失败: {e}")
                    continue
            
            if not questions:
                return [], "未能解析出任何题目"
            
            return questions, ""
            
        except Exception as e:
            return [], f"AI解析失败: {str(e)}"
    
    def parse_questions_from_file(self, file_path: str) -> tuple[List[Question], str]:
        """
        从文件解析题目(支持 Word、Excel、TXT、图片)
        返回: (题目列表, 错误消息)
        """
        path = Path(file_path)
        if not path.exists():
            return [], "文件不存在"
        
        suffix = path.suffix.lower()
        
        # 根据文件类型选择处理方式
        if suffix in ['.txt']:
            return self._parse_txt_file(path)
        elif suffix in ['.doc', '.docx']:
            return self._parse_word_file(path)
        elif suffix in ['.xls', '.xlsx']:
            return self._parse_excel_file(path)
        elif suffix in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
            return self.parse_questions_from_image(str(path))
        else:
            return [], f"不支持的文件格式: {suffix}"
    
    def _parse_txt_file(self, path: Path) -> tuple[List[Question], str]:
        """解析TXT文件"""
        try:
            # 尝试多种编码读取
            content = None
            for encoding in ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']:
                try:
                    content = path.read_text(encoding=encoding)
                    break
                except (UnicodeDecodeError, UnicodeError):
                    continue
            
            if content is None:
                return [], "无法识别文件编码"
            
            return self.parse_questions_from_text(content)
        except Exception as e:
            return [], f"TXT文件解析失败: {str(e)}"
    
    def _parse_word_file(self, path: Path) -> tuple[List[Question], str]:
        """解析Word文档"""
        if not HAS_DOCX:
            return [], "请安装 python-docx 库: pip install python-docx"
        
        try:
            doc = DocxDocument(str(path))
            
            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # 提取表格内容(如果有)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(' | '.join(row_text))
            
            if not paragraphs:
                return [], "Word文档内容为空"
            
            content = '\n'.join(paragraphs)
            return self.parse_questions_from_text(content)
        except Exception as e:
            return [], f"Word文档解析失败: {str(e)}"
    
    def _parse_excel_file(self, path: Path) -> tuple[List[Question], str]:
        """解析Excel文件"""
        if not HAS_OPENPYXL:
            return [], "请安装 openpyxl 库: pip install openpyxl"
        
        try:
            wb = openpyxl.load_workbook(str(path), data_only=True)
            
            all_content = []
            
            for sheet in wb.worksheets:
                sheet_content = []
                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else '' for cell in row]
                    # 过滤空行
                    if any(v.strip() for v in row_values):
                        sheet_content.append(' | '.join(v for v in row_values if v.strip()))
                
                if sheet_content:
                    all_content.extend(sheet_content)
            
            if not all_content:
                return [], "Excel文件内容为空"
            
            content = '\n'.join(all_content)
            return self.parse_questions_from_text(content)
        except Exception as e:
            return [], f"Excel文件解析失败: {str(e)}"
    
    def get_supported_file_types(self) -> Dict[str, List[str]]:
        """获取支持的文件类型"""
        types = {
            'text': ['.txt'],
            'image': ['.png', '.jpg', '.jpeg', '.gif', '.webp']
        }
        
        if HAS_DOCX:
            types['word'] = ['.doc', '.docx']
        
        if HAS_OPENPYXL:
            types['excel'] = ['.xls', '.xlsx']
        
        return types
    
    def get_file_filter_string(self) -> str:
        """获取文件过滤器字符串(用于文件对话框)"""
        filters = []
        types = self.get_supported_file_types()
        
        all_exts = []
        
        if 'text' in types:
            exts = ' '.join(f'*{e}' for e in types['text'])
            filters.append(f"文本文件 ({exts})")
            all_exts.extend(types['text'])
        
        if 'word' in types:
            exts = ' '.join(f'*{e}' for e in types['word'])
            filters.append(f"Word文档 ({exts})")
            all_exts.extend(types['word'])
        
        if 'excel' in types:
            exts = ' '.join(f'*{e}' for e in types['excel'])
            filters.append(f"Excel表格 ({exts})")
            all_exts.extend(types['excel'])
        
        if 'image' in types:
            exts = ' '.join(f'*{e}' for e in types['image'])
            filters.append(f"图片文件 ({exts})")
            all_exts.extend(types['image'])
        
        # 添加"所有支持的格式"
        all_exts_str = ' '.join(f'*{e}' for e in all_exts)
        filters.insert(0, f"所有支持的格式 ({all_exts_str})")
        
        return ';;'.join(filters)
    
    def parse_questions_from_image(self, image_path: str) -> tuple[List[Question], str]:
        """
        从图片解析题目
        返回: (题目列表, 错误消息)
        """
        try:
            # 读取图片并转为base64
            path = Path(image_path)
            if not path.exists():
                return [], "图片文件不存在"
            
            with open(path, 'rb') as f:
                image_data = base64.b64encode(f.read()).decode('utf-8')
            
            # 获取图片格式
            suffix = path.suffix.lower()
            media_type = {
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.gif': 'image/gif',
                '.webp': 'image/webp'
            }.get(suffix, 'image/jpeg')
            
            prompt = """请识别图片中的题目内容，并转换为标准JSON格式。

要求：
1. 识别所有题目
2. 确定题目类型：single(单选), multiple(多选), judge(判断), fill(填空)
3. 提取选项和正确答案(如果图片中有标注)
4. 如果答案不明确，answer字段留空

输出JSON格式：
{
  "questions": [
    {
      "type": "题目类型",
      "question": "题目内容",
      "options": ["A. 内容1", "B. 内容2"...],
      "answer": "答案",
      "explanation": "",
      "difficulty": 3
    }
  ]
}"""
            
            messages = [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{media_type};base64,{image_data}"
                            }
                        },
                        {
                            "type": "text",
                            "text": prompt
                        }
                    ]
                }
            ]
            
            response = self._call_api(messages, use_vision=True)
            data = self._parse_json_response(response)
            
            questions = []
            for q_data in data.get('questions', []):
                try:
                    q = Question(
                        type=q_data.get('type', 'single'),
                        question=q_data.get('question', ''),
                        options=q_data.get('options', []),
                        answer=q_data.get('answer', ''),
                        explanation=q_data.get('explanation', ''),
                        difficulty=q_data.get('difficulty', 3),
                        source='ai_generated'
                    )
                    questions.append(q)
                except Exception as e:
                    print(f"解析单个题目失败: {e}")
                    continue
            
            if not questions:
                return [], "未能从图片中识别出题目"
            
            return questions, ""
            
        except Exception as e:
            return [], f"图片识别失败: {str(e)}"
    
    def generate_questions(self, topic: str, count: int = 5, 
                          type_distribution: str = None,
                          difficulty_range: tuple = (2, 4)) -> tuple[List[Question], str]:
        """
        根据主题生成题目
        返回: (题目列表, 错误消息)
        """
        try:
            if type_distribution is None:
                type_distribution = "单选题、多选题、判断题"
            
            min_diff, max_diff = difficulty_range
            difficulty_str = f"{min_diff}-{max_diff}级"
            
            # 使用 replace 替换占位符，避免内容中的花括号引起 format 错误
            prompt = self.QUESTION_GENERATE_PROMPT
            prompt = prompt.replace("{topic}", topic)
            prompt = prompt.replace("{count}", str(count))
            prompt = prompt.replace("{type_distribution}", type_distribution)
            prompt = prompt.replace("{difficulty_range}", difficulty_str)
            
            messages = [
                {"role": "system", "content": "你是一个专业的出题专家，擅长设计有区分度的考试题目。"},
                {"role": "user", "content": prompt}
            ]
            
            response = self._call_api(messages)
            data = self._parse_json_response(response)
            
            questions = []
            for q_data in data.get('questions', []):
                try:
                    q = Question(
                        type=q_data.get('type', 'single'),
                        question=q_data.get('question', ''),
                        options=q_data.get('options', []),
                        answer=q_data.get('answer', ''),
                        explanation=q_data.get('explanation', ''),
                        difficulty=q_data.get('difficulty', 3),
                        source='ai_generated'
                    )
                    questions.append(q)
                except Exception as e:
                    print(f"生成单个题目失败: {e}")
                    continue
            
            if not questions:
                return [], "未能生成任何题目"
            
            return questions, ""
            
        except Exception as e:
            return [], f"题目生成失败: {str(e)}"
